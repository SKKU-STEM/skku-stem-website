# SKKU Report DOCX — python-docx builder (Pretendard typography, A4 with header/footer chrome)
# 사용: from template import Report
#       r = Report(...); r.cover(...); r.h1(...); r.p(...); r.build('out.docx')
# Pretendard는 viewer 시스템 폰트에 설치되어 있어야 함. 미설치 시 Word가 Calibri/Malgun으로 대체.
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── 디자인 토큰 (SKKU-STEM 사이트 + skku-report-pdf 와 일치) ───
CREAM = 'FAF9F5'
INK = '141413'
INK_70 = '5C5B57'
CORAL = 'CC785C'
CORAL_700 = '8C4D3A'
CREAM_200 = 'F1EFE7'
CREAM_300 = 'E4E1D5'
WARN_YELLOW = 'FFF2CC'

# Pretendard weight 매핑 — python-docx Font.name 으로 노출되는 family 이름
FONT_REG = 'Pretendard'           # weight 400
FONT_MED = 'Pretendard Medium'    # weight 500
FONT_BOLD = 'Pretendard Bold'     # weight 700
FONT_LIGHT = 'Pretendard Light'   # weight 300
FONT_MONO = 'Courier New'

# ─── OOXML 헬퍼 (python-docx가 노출 안 하는 부분 직접 XML 조작) ───
def _set_font_full(run, font_name):
    """ascii / hAnsi / eastAsia / cs 네 가지 슬롯에 동일 폰트 지정 — 한국어 자모도 Pretendard로 렌더."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), font_name)

def _shade(element, hex_color):
    """문단 / 표 셀에 배경색 — pPr 또는 tcPr 자식으로 w:shd 삽입."""
    if hasattr(element, 'get_or_add_pPr'):
        pr = element.get_or_add_pPr()
    elif hasattr(element, 'get_or_add_tcPr'):
        pr = element.get_or_add_tcPr()
    else:
        return
    existing = pr.find(qn('w:shd'))
    if existing is not None:
        pr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#', ''))
    pr.append(shd)

def _page_number_field(run):
    """run 안에 PAGE 필드 삽입 (Word가 자동 갱신)."""
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def _toc_field(run, levels='1-3'):
    """run 안에 TOC 필드 삽입 — viewer가 F9로 갱신."""
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = f'TOC \\o "{levels}" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t'); placeholder.text = '목차는 Word에서 F9를 눌러 갱신하세요.'
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run._r.append(placeholder); run._r.append(fldChar3)

def _add_border(paragraph, color_hex, sides=('top', 'bottom')):
    """문단 상/하/좌/우 보더 — 콜아웃이나 표지 띠에 사용."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), color_hex.replace('#', ''))
        pBdr.append(el)

# ─── Report 클래스 ───
class Report:
    """SKKU-STEM 스타일 보고서 docx 빌더. 메소드를 호출 순서대로 부르면 doc 상태 누적."""

    def __init__(self, doc_title='Report', doc_version='v1.0', doc_footer_url='example.org',
                 author='SKKU-STEM Lab'):
        self.doc = Document()
        self.title = doc_title
        self.version = doc_version
        self.footer_url = doc_footer_url
        self.doc.core_properties.title = doc_title
        self.doc.core_properties.author = author
        self._setup_page()
        self._setup_styles()
        self._setup_header_footer()

    # ─── 초기 설정 ───
    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(22)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)

    def _set_run(self, run, font=FONT_REG, size=10.5, color=INK, bold=False):
        _set_font_full(run, font)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold

    def _setup_styles(self):
        """built-in Heading 1/2/3 스타일을 SKKU 토큰으로 override."""
        styles = self.doc.styles
        # Normal
        normal = styles['Normal']
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(INK)
        # rFonts XML 조작 — Normal 스타일의 ascii + eastAsia 모두 Pretendard로
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rFonts.set(qn(attr), FONT_REG)
        # H1 (Heading 1)
        h1 = styles['Heading 1']
        h1.font.size = Pt(20)
        h1.font.color.rgb = RGBColor.from_string(CORAL_700)
        h1.font.bold = True
        h1_rPr = h1.element.get_or_add_rPr()
        h1_rFonts = h1_rPr.find(qn('w:rFonts'))
        if h1_rFonts is None:
            h1_rFonts = OxmlElement('w:rFonts')
            h1_rPr.append(h1_rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            h1_rFonts.set(qn(attr), FONT_BOLD)
        # H2
        h2 = styles['Heading 2']
        h2.font.size = Pt(15)
        h2.font.color.rgb = RGBColor.from_string(INK)
        h2.font.bold = True
        h2_rPr = h2.element.get_or_add_rPr()
        h2_rFonts = h2_rPr.find(qn('w:rFonts'))
        if h2_rFonts is None:
            h2_rFonts = OxmlElement('w:rFonts')
            h2_rPr.append(h2_rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            h2_rFonts.set(qn(attr), FONT_BOLD)
        # H3
        h3 = styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.color.rgb = RGBColor.from_string(CORAL_700)
        h3.font.bold = True
        h3_rPr = h3.element.get_or_add_rPr()
        h3_rFonts = h3_rPr.find(qn('w:rFonts'))
        if h3_rFonts is None:
            h3_rFonts = OxmlElement('w:rFonts')
            h3_rPr.append(h3_rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            h3_rFonts.set(qn(attr), FONT_BOLD)

    def _setup_header_footer(self):
        section = self.doc.sections[0]
        # 헤더 — 좌:title, 우:version, 가로 라인 아래
        hdr = section.header.paragraphs[0]
        hdr.text = ''
        # 탭 위치: 페이지 폭 = 174mm 글쓰기 영역, 우측 tab을 끝에
        from docx.shared import Cm
        tab_stops = hdr.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Mm(174), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        r1 = hdr.add_run(self.title)
        self._set_run(r1, font=FONT_REG, size=8, color=INK_70)
        r1.add_tab()
        r2 = hdr.add_run(self.version)
        self._set_run(r2, font=FONT_REG, size=8, color=INK_70)
        _add_border(hdr, CREAM_300, sides=('bottom',))
        # 푸터 — 좌:URL, 우:PAGE field
        ftr = section.footer.paragraphs[0]
        ftr.text = ''
        tab_stops_f = ftr.paragraph_format.tab_stops
        tab_stops_f.add_tab_stop(Mm(174), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        rf1 = ftr.add_run(self.footer_url)
        self._set_run(rf1, font=FONT_REG, size=8, color=INK_70)
        rf1.add_tab()
        rf2 = ftr.add_run()
        self._set_run(rf2, font=FONT_REG, size=8, color=INK_70)
        _page_number_field(rf2)
        _add_border(ftr, CREAM_300, sides=('top',))

    # ─── 콘텐츠 헬퍼 ───
    def cover(self, eyebrow, title, subtitle=None, intro=None, footer_meta_lines=None):
        """표지 페이지 — eyebrow + 큰 제목 + 부제 + intro + meta. 끝에 page break."""
        # 상단 cream 띠를 흉내내기 위해 빈 paragraph + shading으로 박스 (간략화: cream 텍스트 박스 1개)
        # docx에서 진짜 페이지 띠 그리기는 까다로워 표지 첫 paragraph 6개를 cream-200 shaded blank으로 처리
        for _ in range(4):
            p = self.doc.add_paragraph()
            _shade(p._p, CREAM_200)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        self.doc.add_paragraph()  # spacer
        # Eyebrow
        p = self.doc.add_paragraph()
        r = p.add_run(eyebrow)
        self._set_run(r, font=FONT_MED, size=11, color=CORAL_700)
        # Title
        p = self.doc.add_paragraph()
        r = p.add_run(title)
        self._set_run(r, font=FONT_BOLD, size=32, color=INK, bold=True)
        p.paragraph_format.space_after = Pt(2)
        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph()
            r = p.add_run(subtitle)
            self._set_run(r, font=FONT_REG, size=14, color=INK_70)
        # spacer
        for _ in range(8):
            self.doc.add_paragraph()
        # Intro
        if intro:
            p = self.doc.add_paragraph()
            r = p.add_run(intro)
            self._set_run(r, font=FONT_REG, size=11, color=INK)
            p.paragraph_format.space_after = Pt(12)
        # Meta lines
        for line in footer_meta_lines or []:
            p = self.doc.add_paragraph()
            r = p.add_run(line)
            self._set_run(r, font=FONT_REG, size=10, color=INK_70)
            p.paragraph_format.space_after = Pt(0)
        self.doc.add_page_break()

    def toc_field(self, heading='목차', levels='1-3'):
        """Word TOC 필드 삽입 — viewer가 Word에서 F9 누르면 갱신됨."""
        self.h1(heading)
        p = self.doc.add_paragraph()
        r = p.add_run()
        self._set_run(r, font=FONT_REG, size=10, color=INK)
        _toc_field(r, levels=levels)
        self.doc.add_page_break()

    def h1(self, text):
        p = self.doc.add_heading(level=1)
        r = p.add_run(text)
        self._set_run(r, font=FONT_BOLD, size=20, color=CORAL_700, bold=True)

    def h2(self, text):
        p = self.doc.add_heading(level=2)
        r = p.add_run(text)
        self._set_run(r, font=FONT_BOLD, size=15, color=INK, bold=True)

    def h3(self, text):
        p = self.doc.add_heading(level=3)
        r = p.add_run(text)
        self._set_run(r, font=FONT_BOLD, size=12, color=CORAL_700, bold=True)

    def p(self, text, bold=False, color=INK):
        """본문 paragraph. text는 단일 문자열 — bold 강조는 인라인 마크업 없이 별도 호출."""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        self._set_run(r, font=FONT_BOLD if bold else FONT_REG, size=10.5, color=color, bold=bold)
        p.paragraph_format.space_after = Pt(6)

    def small(self, text, color=INK_70):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        self._set_run(r, font=FONT_REG, size=9, color=color)

    def bul(self, items):
        """코랄 bullet 리스트. items는 문자열 배열."""
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-4)
            p.paragraph_format.space_after = Pt(2)
            # bullet 문자 (코랄)
            r_bul = p.add_run('• ')
            self._set_run(r_bul, font=FONT_REG, size=10.5, color=CORAL, bold=True)
            r = p.add_run(item)
            self._set_run(r, font=FONT_REG, size=10.5, color=INK)

    def _callout(self, text, prefix, bg):
        p = self.doc.add_paragraph()
        _shade(p._p, bg)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Mm(3)
        p.paragraph_format.right_indent = Mm(3)
        r1 = p.add_run(prefix + ' ')
        self._set_run(r1, font=FONT_BOLD, size=10, color=INK, bold=True)
        r2 = p.add_run(text)
        self._set_run(r2, font=FONT_REG, size=10, color=INK)

    def note(self, text):
        self._callout(text, '※ 참고:', CREAM_200)

    def warn(self, text):
        self._callout(text, '⚠ 주의:', WARN_YELLOW)

    def code(self, text):
        p = self.doc.add_paragraph()
        _shade(p._p, CREAM_200)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Mm(3)
        r = p.add_run(text)
        self._set_run(r, font=FONT_MONO, size=9, color=INK)

    def faq(self, question, answer):
        pq = self.doc.add_paragraph()
        rq = pq.add_run(f'Q. {question}')
        self._set_run(rq, font=FONT_BOLD, size=11, color=CORAL_700, bold=True)
        pq.paragraph_format.space_before = Pt(8)
        pq.paragraph_format.space_after = Pt(2)
        pa = self.doc.add_paragraph()
        ra = pa.add_run(f'A. {answer}')
        self._set_run(ra, font=FONT_REG, size=10, color=INK)
        pa.paragraph_format.space_after = Pt(6)

    def table(self, data, col_widths_mm=None):
        """첫 행 = 헤더(coral-700 배경, cream 텍스트), 나머지 = cream-200 본문."""
        if not data:
            return
        cols = len(data[0])
        tbl = self.doc.add_table(rows=len(data), cols=cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False if col_widths_mm else True
        for r_idx, row in enumerate(data):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.rows[r_idx].cells[c_idx]
                if col_widths_mm:
                    cell.width = Mm(col_widths_mm[c_idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                _shade(cell._tc, CORAL_700 if r_idx == 0 else CREAM_200)
                # 셀 paragraph 0 사용
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(str(cell_text))
                if r_idx == 0:
                    self._set_run(r, font=FONT_BOLD, size=9, color=CREAM, bold=True)
                else:
                    self._set_run(r, font=FONT_REG, size=9, color=INK)
        # 본문 paragraph 위/아래 여백
        for spacer_pos in (0, 1):
            pass
        # 표 뒤 빈 paragraph 1개로 다음 내용과 간격 확보
        gap = self.doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(4)
        return tbl

    def page_break(self):
        self.doc.add_page_break()

    def image(self, path, width_mm=None):
        kwargs = {'width': Mm(width_mm)} if width_mm else {}
        self.doc.add_picture(str(path), **kwargs)

    def build(self, out_path):
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(out_path)
        return out_path
